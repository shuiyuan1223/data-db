"""Convert markdown files to docx with basic formatting."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = Path(__file__).parent

MD_FILES = [
    "测试策略与测试设计文档.md",
    "模型性能测试报告.md",
    "优化分析文档.md",
]


def parse_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        # skip separator rows like |---|---|
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        if i == 0:
                            run.bold = True


def md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    table_buffer = []
    in_table = False
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.size = Pt(8)
                run.font.name = 'Consolas'
                p.paragraph_format.left_indent = Inches(0.3)
                code_buffer = []
                in_code_block = False
            else:
                # Flush table if any
                if in_table:
                    add_table(doc, parse_table(table_buffer))
                    table_buffer = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table lines
        if line.strip().startswith('|'):
            in_table = True
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            add_table(doc, parse_table(table_buffer))
            table_buffer = []
            in_table = False

        # Skip horizontal rules
        if line.strip() == '---' or line.strip() == '***':
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).replace('**', '').replace('*', '')
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith('>'):
            text = line.strip().lstrip('> ').replace('**', '').replace('*', '')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # List items
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            # Handle bold within text
            p = doc.add_paragraph(style='List Bullet' if list_match.group(2) in ['-', '*', '+'] else 'List Number')
            # Simple bold handling
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.bold = True
            i += 1
            continue

        # Regular paragraph with bold handling
        text = line.strip()
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(10.5)
            if idx % 2 == 1:
                run.bold = True
        i += 1

    # Flush remaining table
    if in_table:
        add_table(doc, parse_table(table_buffer))

    doc.save(str(docx_path))
    print(f"  ✅ {docx_path.name}")


if __name__ == '__main__':
    for md_name in MD_FILES:
        md_path = OUTPUT_DIR / md_name
        docx_path = OUTPUT_DIR / md_name.replace('.md', '.docx')
        if md_path.exists():
            md_to_docx(md_path, docx_path)
        else:
            print(f"  ⚠️ {md_name} not found")
    print("Done.")
